from flask import Flask, request, render_template, send_file
import re
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO

app = Flask(__name__)

def extract_values_from_txt(file_content, patterns):
    encodings = ['utf-8', 'cp1251', 'latin-1']
    for encoding in encodings:
        try:
            content = file_content.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        raise ValueError("Не удалось открыть файл с доступными кодировками")

    values = {}
    for key, pattern in patterns.items():
        match = re.search(pattern, content)
        if match:
            values[key] = match.groups()
        else:
            values[key] = (None, None, None, None)
    
    return values

def create_document_with_values(values, gender):
    doc = Document()

    title = doc.add_heading(level=1)
    title_run = title.add_run('Заключение')
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.name = 'Arial'
    title_run.font.color.rgb = RGBColor(0, 0, 0)

# Шаблонный текст с метками
    if gender.lower() == 'мужчина':
        template_text = """
        Соотношение челюстей в сагиттальной плоскости.
        Межапикальный угол (<ANB) {ANB}˚, что соответствует соотношению челюстей {ANB_class} (N = 2,0˚ ± 2,0˚).
        Угол Бета (< Beta Angle) – {Beta_Angle}˚, что соответствует соотношению челюстей {Beta_Angle_class} (N = 31,0˚ ± 4,0˚). 
        Параметр Wits (Wits Appraisal) – {Wits_Appraisal} мм, что указывает на {Wits_Appraisal_class_disproportion} диспропорции в расположении апикальных базисов верхней и нижней челюстей в сагиттальной плоскости
        и соответствует {Wits_Appraisal_class} (N (М) = -1,1 мм ± 2,0 мм).
        Соотношение челюстей в сагиттальной плоскости по методике Sassouni (B to A Point Arc) соответствует {B_to_A_Point_Arc_class} — 
        базальная дуга проходит на {B_to_A_Point_Arc} мм {B_to_A_Point_Arc_direction} от точки В (N = 0,0 мм ± 3,0 мм).
        Параметр APDI, указывающий на дисплазию развития челюстей в сагиттальной плоскости, составляет {APDI}˚ и соответствует {Apdi_class} (N = 81,4˚ ± 5,0˚).

        Размер и положение верхней челюсти.
        Размер основания верхней челюсти (PNS-A) – {PNS_A} мм, что соответствует {Pns_sn_value_result} (N = {SN_result} мм ±  3,5 мм). 
        Положение верхней челюсти по сагиттали  (<SNA) – {SNA}˚, что {SNA_class} (N = 82,0˚ ±  3,0˚).
        Положение верхней челюсти по вертикали  (<SN-Palatal Plane) – {SN_Palatal_Plane}˚, что {SN_palatal_plane_class} (N= 8,0˚ ± 3,0˚).

        Размер и положение нижней челюсти.
        Размер тела нижней челюсти (Go-Me) – {Go_Me} мм, что {Go_me_value_result} (N = {Go_me_result} мм ±  5,0 мм).
        Положение нижней челюсти по сагиттали  (<SNB) – {SNB}˚, что {SNB_class} (N = 80,0˚ ± 3,0˚). 
        Положение нижней челюсти по вертикали (<MP-SN) – {MP_SN}˚, что {MP_SN_class} (N= 32,0˚ ± 4,0˚).

        Вертикальные параметры.
        Гониальный угол (<Ar-Go-Me) – {Ar_Go_Me}˚, что {Ar_Go_Me_class} (N = 130,0˚ ±  5,0˚ ).
        Верхний гониальный угол (<Ar-Go-N) – {Ar_Go_Na}˚, что {Ar_Go_Na_class} (N = 55,0˚ ±  3,0˚ ).
        Нижний гониальный угол равен (<N-Go-Me) – {Na_Go_Me}˚, что {Na_Go_Me_class} (N = 75,0˚ ±  3,0˚).     
        Вертикальное лицевое соотношение (N-ANS/ANS-Gn) {ANS_result_class} –  {ANS_result}  (N = 0,8 ±  0,09).
        Отношение задней высоты лица к передней (S-Go/N-Gn) – {SGo_NGn}%, что {Sgo_ngn_class} (N = 63,0% ±  2,0%).
        Высота нижней трети лица по Ricketts (<ANS-Xi-Pm) – {ANS_Xi_Pm}˚, что {I_VP_result_class} (N = IVP {I_VP_result}˚ ± 5,5˚).
        Параметр ODI – {ODI}˚, что {ODI_class} (N = 74,5˚ ±  5,0˚).
        
        Оценка положения резцов и моляров.
        Межрезцовый угол ( U1\L1) – {U1_L1}˚, что {U1_L1_class} ( N= 130,0˚ ± 6,0˚).
        {Palatal_Plane_class}. {IMPA_class}.
        Наклон резцов на верхней \ нижней челюсти в норме \ тенденция к протрузии \ ретрузии резцов на верхней\ нижней челюсти.
        {Uagh_Uahh_Updn_Lpdn_result}

        Оценка мягких тканей лица.
        Профиль (<gl-sn-pog) {G_SN_PO_class}. 
        Носогубный угол  (<col-sn-UL) – {Col_Sn_UL}˚, что соответствует {Сol_sn_ul_class} (N = 102,0˚ ± 8,0˚). 
        Положение верхней губы – {Upper_lip} мм (N = -4,0 мм ± 2,0 мм). Положение нижней губы – {Lower_lip} мм ( N = -2,0 мм ±2,0 мм).

        """
    else:  # Предположим, что если не мужчина, то женщина
        template_text = """
        Соотношение челюстей в сагиттальной плоскости.
        Межапикальный угол (<ANB) {ANB}˚, что соответствует соотношению челюстей {ANB_class} (N = 2,0˚ ± 2,0˚).
        Угол Бета (< Beta Angle) – {Beta_Angle}˚, что соответствует соотношению челюстей {Beta_Angle_class} (N = 31,0˚ ± 4,0˚). 
        Параметр Wits (Wits Appraisal) – {Wits_Appraisal} мм, что указывает на {Wits_Appraisal_class_disproportion} диспропорции в расположении апикальных базисов верхней и нижней челюстей в сагиттальной плоскости
        и соответствует {Wits_appraisal_woman_class} (N (Ж) = -0,4 мм ± 2,5 мм).
        Соотношение челюстей в сагиттальной плоскости по методике Sassouni (B to A Point Arc) соответствует {B_to_A_Point_Arc_class} — 
        базальная дуга проходит на {B_to_A_Point_Arc} мм {B_to_A_Point_Arc_direction} от точки В (N = 0,0 мм ± 3,0 мм).
        Параметр APDI, указывающий на дисплазию развития челюстей в сагиттальной плоскости, составляет {APDI}˚ и соответствует {Apdi_class} (N = 81,4˚ ± 5,0˚).

        Размер и положение верхней челюсти.
        Размер основания верхней челюсти (PNS-A) – {PNS_A} мм, что соответствует {Pns_sn_value_result} (N = {SN_result} мм ±  3,5 мм). 
        Положение верхней челюсти по сагиттали  (<SNA) – {SNA}˚, что {SNA_class} (N = 82,0˚ ±  3,0˚).
        Положение верхней челюсти по вертикали  (<SN-Palatal Plane) – {SN_Palatal_Plane}˚, что {SN_palatal_plane_class} (N= 8,0˚ ± 3,0˚).

        Размер и положение нижней челюсти.
        Размер тела нижней челюсти (Go-Me) – {Go_Me} мм, что {Go_me_value_result} (N = {Go_me_result} мм ±  5,0 мм).
        Положение нижней челюсти по сагиттали  (<SNB) – {SNB}˚, что {SNB_class} (N = 80,0˚ ± 3,0˚). 
        Положение нижней челюсти по вертикали (<MP-SN) – {MP_SN}˚, что {MP_SN_class} (N= 32,0˚ ± 4,0˚).

        Вертикальные параметры.
        Гониальный угол (<Ar-Go-Me) – {Ar_Go_Me}˚, что {Ar_Go_Me_class} (N = 130,0˚ ±  5,0˚ ).
        Верхний гониальный угол (<Ar-Go-N) – {Ar_Go_Na}˚, что {Ar_Go_Na_class} (N = 55,0˚ ±  3,0˚ ).
        Нижний гониальный угол равен (<N-Go-Me) – {Na_Go_Me}˚, что {Na_Go_Me_class} (N = 75,0˚ ±  3,0˚).     
        Вертикальное лицевое соотношение (N-ANS/ANS-Gn) {ANS_result_class} –  {ANS_result}  (N = 0,8 ±  0,09).
        Отношение задней высоты лица к передней (S-Go/N-Gn) – {SGo_NGn}%, что {Sgo_ngn_class} (N = 63,0% ±  2,0%).
        Высота нижней трети лица по Ricketts (<ANS-Xi-Pm) – {ANS_Xi_Pm}˚, что {I_VP_result_class} (N = IVP {I_VP_result}˚ ± 5,5˚).
        Параметр ODI – {ODI}˚, что {ODI_class} (N = 74,5˚ ±  5,0˚).
        
        Оценка положения резцов и моляров.
        Межрезцовый угол ( U1\L1) – {U1_L1}˚, что {U1_L1_class} ( N= 130,0˚ ± 6,0˚).
        {Palatal_Plane_class}. {IMPA_class}.
        Наклон резцов на верхней \ нижней челюсти в норме \ тенденция к протрузии \ ретрузии резцов на верхней\ нижней челюсти.
        {Uagh_Uahh_Updn_Lpdn_result}

        Оценка мягких тканей лица.
        Профиль (<gl-sn-pog) {G_SN_PO_class}. 
        Носогубный угол  (<col-sn-UL) – {Col_Sn_UL}˚, что соответствует {Сol_sn_ul_class} (N = 102,0˚ ± 8,0˚). 
        Положение верхней губы – {Upper_lip} мм (N = -4,0 мм ± 2,0 мм). Положение нижней губы – {Lower_lip} мм ( N = -2,0 мм ±2,0 мм).
        """

    # Определение значения класса по ANB
    anb_value = float(values['ANB'][0]) if values['ANB'][0] else None
    if anb_value is not None:
        anb_class = (
            "по I скелетному классу с тенденцией к II классу" if 3.8 <= anb_value <= 4.0 else
            "по I скелетному классу с тенденцией к III классу" if 0 <= anb_value <= 0.2 else
            "по III скелетному классу с тенденцией к I классу" if -0.2 <= anb_value < 0 else
            "по II скелетному классу с тенденцией к I классу" if 4.0 <= anb_value <= 4.2 else
            "по I скелетному классу" if 0 < anb_value < 3.8 else
            "по II скелетному классу" if -4.0 < anb_value < 0 else
            "по III скелетному классу" if anb_value > 4.0 else
            "по II скелетному классу" if anb_value < -4.0 else
            ""
        )
    else:
        anb_class = 'значение не найдено'

    # Определение значения класса по Beta Angle
    beta_angle_value = float(values['Beta_Angle'][0]) if values['Beta_Angle'][0] else None
    if beta_angle_value is not None:
        beta_angle_class = (
            "по I скелетному классу с тенденцией к II классу" if 27.0 <= beta_angle_value <= 27.2 else
            "по I скелетному классу с тенденцией к III классу" if 34.8 <= beta_angle_value <= 35.0 else
            "по III скелетному классу с тенденцией к I классу" if 35.1 <= beta_angle_value <= 35.2 else
            "по II скелетному классу с тенденцией к I классу" if 26.8 <= beta_angle_value <= 26.9 else
            "по I скелетному классу" if 27.0 < beta_angle_value < 35.0 else
            "по II скелетному классу" if beta_angle_value < 27.0 else
            "по III скелетному классу" if beta_angle_value > 35.0 else
            ""
        )
    else:
        beta_angle_class = 'значение не найдено'
    
    # Определение значения класса по Wits Appraisal man
    wits_appraisal_value = float(values['Wits_Appraisal'][0]) if values['Wits_Appraisal'][0] else None
    if wits_appraisal_value is not None:
        wits_appraisal_class = (
            "по I скелетному классу с тенденцией к II классу" if 0.8 <= wits_appraisal_value <= 1.0 else
            "по I скелетному классу с тенденцией к III классу" if -3.0 <= wits_appraisal_value <= -2.8 else
            "по III скелетному классу с тенденцией к I классу" if -3.5 <= wits_appraisal_value <= -3.3 else
            "по II скелетному классу с тенденцией к I классу" if 1.2 <= wits_appraisal_value <= 1.4 else
            "по I скелетному классу" if -3.1 <= wits_appraisal_value <= 1.1 else
            "по II скелетному классу" if wits_appraisal_value > 1.2 else
            "по III скелетному классу" if wits_appraisal_value < -3.2 else
            ""
        )
    else:
        wits_appraisal_class = 'значение не найдено'

        # Определение значения класса по Wits Appraisal woman
    wits_appraisal_woman_value = float(values['Wits_Appraisal'][0]) if values['Wits_Appraisal'][0] else None
    if wits_appraisal_woman_value is not None:
        wits_appraisal_woman_class = (
            "по I скелетному классу с тенденцией к II классу" if 1.9 <= wits_appraisal_woman_value <= 2.0 else
            "по I скелетному классу с тенденцией к III классу" if -2.8 <= wits_appraisal_woman_value <= -2.6 else
            "по III скелетному классу с тенденцией к I классу" if -3.2 <= wits_appraisal_woman_value <= -3.0 else
            "по II скелетному классу с тенденцией к I классу" if 2.2 <= wits_appraisal_woman_value <= 2.4 else
            "по I скелетному классу" if -2.9 <= wits_appraisal_woman_value <= 2.1 else
            "по II скелетному классу" if wits_appraisal_woman_value > 2.2 else
            "по III скелетному классу" if wits_appraisal_woman_value < -3.0 else
            ""
        )
    else:
        wits_appraisal_woman_class = 'значение не найдено'

    # Определение наличия или отсутствия диспропорции по Wits Appraisal
    if wits_appraisal_value is not None:
        wits_appraisal_class_disproportion = (
            "отсутствие" if -0.4 <= wits_appraisal_value <= 1.1 else
            "наличие"
        )
    else:
        wits_appraisal_class_disproportion = 'значение не найдено'

    # Определение значения класса по B_to_A_Point_Arc
    b_to_a_point_arc_value = float(values['B_to_A_Point_Arc'][0]) if values['B_to_A_Point_Arc'][0] else None
    if b_to_a_point_arc_value is not None:
        b_to_a_point_arc_class = (
            "I классу с тенденцией к II классу" if -2.7 <= b_to_a_point_arc_value <= -2.9 else
            "I классу с тенденцией к III классу" if 2.7 <= b_to_a_point_arc_value <= 2.9 else
            "III классу с тенденцией к I классу" if 3.1 <= b_to_a_point_arc_value <= 3.3 else
            "II классу с тенденцией к I классу" if -3.3 <= b_to_a_point_arc_value <= -3.1 else
            "I классу" if -3.0 <= b_to_a_point_arc_value <= 3.0 else
            "II классу" if b_to_a_point_arc_value < -3.0 else
            "III классу" if b_to_a_point_arc_value > 3.0 else
            ""
        )
        # Определение значения кпереди/кзади
        b_to_a_point_arc_direction = (
            "кпереди" if b_to_a_point_arc_value > 0 else
            "кзади" if b_to_a_point_arc_value < 0 else
            "значение 0"
        )
    else:
        b_to_a_point_arc_class = 'значение не найдено'
        b_to_a_point_arc_direction = 'значение не найдено'

    # Определение значения класса по APDI
    apdi_value = float(values['APDI'][0]) if values['APDI'][0] else None
    if apdi_value is not None:
        apdi_class = (
            "I классу с тенденцией к II классу" if 76.5 <= apdi_value <= 76.7 else
            "I классу с тенденцией к III классу" if 86.1 <= apdi_value <= 86.3 else
            "III классу с тенденцией к I классу" if 86.5 <= apdi_value <= 86.7 else
            "II классу с тенденцией к I классу" if 76.1 <= apdi_value <= 76.3 else
            "I классу" if 76.4 <= apdi_value <= 86.4 else
            "II классу" if apdi_value < 76.4 else
            "III классу" if apdi_value > 86.4 else
            ""
    )
    else:
        apdi_class = 'значение не найдено'

    # Определение нормы N параметра PNS-A, Go-Me
    sn_value = float(values['SN'][0]) if values['SN'][0] else None
    sn_result = round(sn_value * 0.7, 1) if sn_value is not None else 'значение не найдено'
    go_me_result = round(sn_value * 21 / 20, 1) if sn_value is not None else 'значение не найдено'
    
    # Тернарное выражение для определения соответствующего текста
    pns_sn_value = float(values['PNS_A'][0]) if values['PNS_A'][0] else None
    if pns_sn_value is not None:
        pns_sn_value_result = (
        "соответствует индивидуальной норме" if abs(sn_result - pns_sn_value) <= 3.5 else
        "соответствует уменьшению" if sn_result < pns_sn_value - 3.5 else
        "соответствует увеличению" if sn_result > pns_sn_value + 3.5 else
        "значение не определено"  # В случае, если ни одно из условий не выполнено
    )
    else:
        pns_sn_value_result = "значение PNS_A не найдено"

        # Тернарное выражение для определения соответствующего текста Go-Me
    go_me_value = float(values['Go_Me'][0]) if values['Go_Me'][0] else None

    if go_me_value is not None:
        go_me_value_result = (
            "соответствует индивидуальной норме" if abs(go_me_result - go_me_value) <= 5.0 else
            "соответствует уменьшению" if go_me_result < go_me_value - 5.0 else
            "соответствует увеличению" if go_me_result > go_me_value + 5.0 else
            "значение не определено"  # В случае, если ни одно из условий не выполнено
    )
    else:
        go_me_value_result = "значение Go_Me не найдено"

        # Определение значения класса по SNA
    sna_value = float(values['SNA'][0]) if values['SNA'][0] else None
    if sna_value is not None:
        sna_class = (
            "соответствует норме с тенденцией к ретрогнатии" if 79.1 <= sna_value <= 79.3 else
            "соответствует норме с тенденцией к прогнатии" if 84.7 <= sna_value <= 84.9 else
            "соответствует прогнатии с тенденцией к норме" if 85.1 <= sna_value <= 85.3 else
            "соответствует ретрогнатии с тенденцией к норме" if 78.7 <= sna_value <= 78.9 else
            "соответствует норме" if 79.0 <= sna_value <= 85.0 else
            "соответствует ретрогнатии" if sna_value < 79.0 else
            "соответствует прогнатии" if sna_value > 85.0 else
            ""
    )
    else:
        sna_class = 'значение не найдено'

        # Определение значения класса по SNB
    snb_value = float(values['SNB'][0]) if values['SNB'][0] else None
    if snb_value is not None:
        snb_class = (
            "соответствует норме с тенденцией к ретрогнатии" if 77.1 <= snb_value <= 77.3 else
            "соответствует норме с тенденцией к прогнатии" if 82.7 <= snb_value <= 82.9 else
            "соответствует прогнатии с тенденцией к норме" if 83.1 <= snb_value <= 83.3 else
            "соответствует ретрогнатии с тенденцией к норме" if 76.7 <= snb_value <= 76.9 else
            "соответствует норме" if 77.0 <= snb_value <= 83.0 else
            "соответствует ретрогнатии" if snb_value < 77.0 else
            "соответствует прогнатии" if snb_value > 83.0 else
            ""
    )
    else:
        snb_class = 'значение не найдено'


    # Определение значения класса по sn_palatal_plane
    sn_palatal_plane_value = float(values['SN_Palatal_Plane'][0]) if values['SN_Palatal_Plane'][0] else None
    if sn_palatal_plane_value is not None:
        sn_palatal_plane_class = (
            "соответствует нормоинкликации" if 5.0 <= sn_palatal_plane_value <= 11.0 else
            "соответствует антеинкликации" if sn_palatal_plane_value < 5.0 else
            "соответствует ретроинкликации" if sn_palatal_plane_value > 11.0 else
            "соответствует антеинкликации с тенденцией к нормоинкликации" if 4.7 <= sn_palatal_plane_value <= 4.9 else
            "соответствует нормоинкликации с тенденцией к антеинкликации" if 5.1 <= sn_palatal_plane_value <= 5.3 else
            "соответствует нормоинкликации с тенденцией к ретроинкликации" if 10.7 <= sn_palatal_plane_value <= 10.9 else
            "соответствует ретроинкликации с тенденцией к нормоинкликации" if 11.1 <= sn_palatal_plane_value <= 11.3 else
            ""
    )
    else:
        sn_palatal_plane_class = 'значение не найдено'

    # Определение значения класса по sn_palatal_plane
    mp_sn_value = float(values['MP_SN'][0]) if values['MP_SN'][0] else None
    if mp_sn_value is not None:
        mp_sn_class = (
            "соответствует нормоинкликации" if 28.0 <= mp_sn_value <= 36.0 else
            "соответствует антеинкликации" if mp_sn_value < 28.0 else
            "соответствует ретроинкликации" if mp_sn_value > 36.0 else
            "соответствует антеинкликации с тенденцией к нормоинкликации" if 27.7 <= mp_sn_value <= 27.9 else
            "соответствует нормоинкликации с тенденцией к антеинкликации" if 28.1 <= mp_sn_value <= 28.3 else
            "соответствует нормоинкликации с тенденцией к ретроинкликации" if 35.7 <= mp_sn_value <= 35.9 else
            "соответствует ретроинкликации с тенденцией к нормоинкликации" if 36.1 <= mp_sn_value <= 36.3 else
            ""
    )
    else:
        mp_sn_class = 'значение не найдено'

    # Определение значения класса по ar_go_me
    ar_go_me_value = float(values['Ar_Go_Me'][0]) if values['Ar_Go_Me'][0] else None
    ar_go_me_class = (
    'соответствует норме' if 125.0 <= ar_go_me_value <= 135.0 else
    'соответствует уменьшению' if ar_go_me_value < 125.0 else
    'соответствует увеличению' if ar_go_me_value > 135.0 else
    ''
    ) if ar_go_me_value is not None else 'значение не найдено'

    # Определение значения класса по ar_go_na
    ar_go_na_value = float(values['Ar_Go_Na'][0]) if values['Ar_Go_Na'][0] else None

    if ar_go_na_value is not None:
        ar_go_na_class = (
            "соответствует норме" if 52.0 <= ar_go_na_value <= 58.0 else
            "соответствует уменьшению" if ar_go_na_value < 52.0 else
            "соответствует увеличению" if ar_go_na_value > 58.0 else
            ""
    )
    else:
        ar_go_na_class = 'значение не найдено'

    # Определение значения класса по na_go_me
    na_go_me_value = float(values['Na_Go_Me'][0]) if values['Na_Go_Me'][0] else None

    na_go_me_class = (
        "соответствует норме" if 72.0 <= na_go_me_value <= 78.0 
        else "соответствует уменьшению" if na_go_me_value < 72.0 
        else "соответствует увеличению" if na_go_me_value > 78.0 
        else "значение не найдено"
    )

    # Определение нормы параметра N_ANS, ANS_Gn
    n_ans_value = float(values['N_ANS'][0]) if values['N_ANS'][0] else None
    ans_gn_value = float(values['ANS_Gn'][0]) if values['ANS_Gn'][0] else None
    ans_result = round(n_ans_value / ans_gn_value, 1) if sn_value is not None else 'значение не найдено'

    # Определение значения класса по ans
    ans_result_class = (
    "гармоничное" if ans_result is not None and 0.71 <= ans_result <= 0.89 else
    "негармоничное"
    ) if ans_result is not None else 'значение не найдено'

    # Определение значения класса по sgo_ngn
    sgo_ngn_value = float(values['SGo_NGn'][0]) if values['SGo_NGn'][0] else None
    sgo_ngn_class = (
        "соответствует нейтральному типу роста" if 61.0 <= sgo_ngn_value <= 65.0 else
        "соответствует горизонтальному типу роста" if sgo_ngn_value > 65.0 else
        "соответствует вертикальному типу роста" if sgo_ngn_value < 61.0 else
        "значение не найдено"
    ) if sgo_ngn_value is not None else 'значение не найдено'

    # Определение нормы параметра I_VP
    fma_value = float(values['FMA'][0]) if values['FMA'][0] else None
    nababtgn_value = float(values['NaBa_PtGn'][0]) if values['NaBa_PtGn'][0] else None
    i_vp_result = round(58.0 + fma_value * 0.2 - nababtgn_value * 0.2, 1) if sn_value is not None else 'значение не найдено'

    # Установленный диапазон и допустимое отклонение
    established_value = i_vp_result  # Пример установленного значения
    tolerance = 5.5

    # Вычисляем границы диапазона
    lower_bound = established_value - tolerance
    upper_bound = established_value + tolerance

    # Тернарное выражение для определения результата
    i_vp_result_class = (
    "соответствует норме" 
    if i_vp_result is not None and lower_bound <= i_vp_result <= upper_bound 
    else "соответствует уменьшению" 
    if i_vp_result is not None and i_vp_result < lower_bound 
    else "соответствует увеличению" 
    if i_vp_result is not None and i_vp_result > upper_bound 
    else "значение не найдено"
    )

    # Определение значения класса по sgo_ngn
    odi_value = float(values['ODI'][0]) if values['ODI'][0] else None
    odi_class = (
        "соответствует норме" if 69.5 <= odi_value <= 79.5 else
        "соответствует вертикальной дизокклюзии" if odi_value < 69.5 else
        "соответствует глубокой резцовой окклюзии" if odi_value > 79.5 else
        "соответствует норме с тенденцией к вертикальной дизокклюзии" if 69.6 <= odi_value <= 69.8 else
        "соответствует норме с тенденцией к глубокой резцовой окклюзии" if 79.2 <= odi_value <= 79.4 else
        "соответствует вертикальной дизокклюзии с тенденцией к норме" if 69.2 <= odi_value <= 69.4 else
        "соответствует глубокой резцовой окклюзии с тенденцией к норме" if 79.6 <= odi_value <= 79.8 else
        "значение не найдено"
    )

    # Определение значения класса по u1_l1
    u1_l1_value = float(values['U1_L1'][0]) if values.get('U1_L1') and values['U1_L1'][0] else None

    # Теперь используем тернарное выражение с учетом, что u1_l1_value может быть None
    u1_l1_class = (
        "значение не найдено" if u1_l1_value is None else
        "соответствует норме" if 124.0 <= u1_l1_value <= 136.0 else
        "соответствует уменьшению" if u1_l1_value < 124.0 else
        "соответствует увеличению" if u1_l1_value > 136.0 else
        "соответствует норме с тенденцией к уменьшению" if 123.8 <= u1_l1_value <= 123.9 else
        "соответствует норме с тенденцией к увеличению" if 136.1 <= u1_l1_value <= 136.2 else
        "соответствует уменьшению с тенденцией к норме" if 123.6 <= u1_l1_value <= 123.7 else
        "соответствует увеличению с тенденцией к норме" if 136.3 <= u1_l1_value <= 136.4 else
        "значение не найдено"
    )

     # Определение значения класса по Palatal_Plane
    palatal_plane_value = float(values['Palatal_Plane'][0]) if values.get('Palatal_Plane') and values['Palatal_Plane'][0] else None

    # Используем проверку на None и тернарное выражение для установки класса
    if palatal_plane_value is not None:
        palatal_plane_class = (
            "Положение резцов на верхней челюсти соответствует норме" if 105.0 <= palatal_plane_value <= 115.0 else
            f"Ретрузия резцов на верхней челюсти на ({105.0 - palatal_plane_value:.1f})˚" if palatal_plane_value < 105.0 else
            f"Протрузия резцов на верхней челюсти на ({palatal_plane_value - 115.0:.1f})˚" if palatal_plane_value > 115.0 else
            ""
    )
    else:
        palatal_plane_class = 'значение не найдено'


    # Определение значения класса по impa
    impa_value = float(values['IMPA'][0]) if values.get('IMPA') and values['IMPA'][0] else None

    # Используем проверку на None и тернарное выражение для установки класса
    if impa_value is not None:
        impa_class = (
            "Положение резцов на нижней челюсти соответствует норме" if 90.0 <= impa_value <= 100.0 else
            f"Ретрузия резцов на нижней челюсти на {90.0 - impa_value:.1f}˚" if impa_value < 90.0 else
            f"Протрузия резцов на нижней челюсти на {impa_value - 100.0:.1f}˚" if impa_value > 100.0 else
            ""
    )
    else:
        impa_class = 'значение не найдено'


    # Преобразование значений в числа с плавающей точкой, если они присутствуют
    UADH = float(values['UADH'][0]) if values.get('UADH') and values['UADH'][0] else None
    UAHH = float(values['UAHH'][0]) if values.get('UAHH') and values['UAHH'][0] else None
    UPDH = float(values['UPDH'][0]) if values.get('UPDH') and values['UPDH'][0] else None
    LPDH = float(values['LPDH'][0]) if values.get('LPDH') and values['LPDH'][0] else None

    # Определение класса по значениям UADH, UAHH, UPDH и LPDH
    if UADH is not None and UAHH is not None and UPDH is not None and LPDH is not None:
        uagh_uahh_updn_lpdn_result = (
            "Зубоальвеолярное укорочение в области резцов на верхней (U1-PP) и нижней (L1-MP) челюстях, моляров на верхней (U6-PP) и нижней (L6-MP) челюстях" if UADH < 25.0 and UAHH < 38.0 and UPDH < 21.0 and LPDH < 29.0 else
            "Зубоальвеолярное укорочение в области резцов на верхней (U1-PP) и нижней (L1-MP) челюстях, моляров на верхней (U6-PP) челюсти" if UADH < 25.0 and UAHH < 38.0 and UPDH < 21.0 and 29.0 <= LPDH <= 33.0 else
            "Зубоальвеолярное укорочение в области резцов на верхней (U1-PP) и нижней (L1-MP) челюстях, моляров нижней (L6-MP) челюсти" if UADH < 25.0 and UAHH < 38.0 and 21.0 <= UPDH <= 24.0 and LPDH < 29.0 else
            "Зубоальвеолярное укорочение в области резцов на верхней (U1-PP) и нижней (L1-MP) челюстях" if UADH < 25.0 and UAHH < 38.0 and 21.0 <= UPDH <= 24.0 and 29.0 <= LPDH <= 33.0 else
            "Зубоальвеолярное укорочение в области резцов на верхней (U1-PP) челюсти, моляров на верхней (U6-PP) и нижней (L6-MP) челюстях" if UADH < 25.0 and 38.0 <= UAHH <= 42.0 and UPDH < 21.0 and LPDH < 29.0 else
            "Зубоальвеолярное укорочение в области резцов на верхней (U1-PP) челюсти" if UADH < 25.0 and 38.0 <= UAHH <= 42.0 and UPDH < 21.0 and 29.0 <= LPDH <= 33.0 else
            "Зубоальвеолярное укорочение в области резцов на верхней (U1-PP) челюсти, моляров нижней (L6-MP) челюсти" if UADH < 25.0 and 38.0 <= UAHH <= 42.0 and 21.0 <= UPDH <= 24.0 and LPDH < 29.0 else
            "Зубоальвеолярное укорочение в области резцов на верхней (U1-PP) челюсти" if UADH < 25.0 and 38.0 <= UAHH <= 42.0 and 21.0 <= UPDH <= 24.0 and 29.0 <= LPDH <= 33.0 else
            "Зубоальвеолярное укорочение в области резцов на верхней (U1-PP) челюсти, зубоальвеолярное удлинение в области резцов на нижней (L1-MP) челюсти, зубоальвеолярное укорочение моляров на верхней (U6-PP) и нижней (L6-MP) челюстях" if UADH < 25.0 and UAHH > 42.0 and UPDH < 21.0 and LPDH < 29.0 else
            "Зубоальвеолярное укорочение в области резцов на верхней (U1-PP) челюсти, зубоальвеолярное удлинение в области резцов на нижней (L1-MP) челюсти, зубоальвеолярное укорочение моляров на верхней (U6-PP) челюсти" if UADH < 25.0 and UAHH > 42.0 and UPDH < 21.0 and 29.0 <= LPDH <= 33.0 else
            "Зубоальвеолярное укорочение в области резцов на верхней (U1-PP) челюсти, зубоальвеолярное удлинение в области резцов на нижней (L1-MP) челюсти, зубоальвеолярное укорочение моляров на нижней (L6-MP) челюсти" if UADH < 25.0 and UAHH > 42.0 and 21.0 <= UPDH <= 24.0 and LPDH < 29.0 else
            "Зубоальвеолярное укорочение в области резцов на верхней (U1-PP) челюсти, зубоальвеолярное удлинение в области резцов на нижней (L1-MP) челюсти" if UADH < 25.0 and UAHH > 42.0 and 21.0 <= UPDH <= 24.0 and 29.0 <= LPDH <= 33.0 else
            "Зубоальвеолярное укорочение в области резцов на нижней (L1-MP) челюсти, зубоальвеолярное укорочение моляров на верхней (U6-PP) и нижней (L6-MP) челюстях" if 25.0 <= UADH <= 31.0 and UAHH < 38.0 and UPDH < 21.0 and LPDH < 29.0 else
            "Зубоальвеолярное укорочение в области резцов на нижней (L1-MP) челюсти" if 25.0 <= UADH <= 31.0 and UAHH < 38.0 and UPDH < 21.0 and 29.0 <= LPDH <= 33.0 else
            "Зубоальвеолярное укорочение в области резцов на нижней (L1-MP) челюсти, зубоальвеолярное укорочение моляров на нижней (L6-MP) челюсти" if 25.0 <= UADH <= 31.0 and UAHH < 38.0 and 21.0 <= UPDH <= 24.0 and LPDH < 29.0 else
            "Зубоальвеолярное укорочение в области резцов на нижней (L1-MP) челюсти" if 25.0 <= UADH <= 31.0 and UAHH < 38.0 and 21.0 <= UPDH <= 24.0 and 29.0 <= LPDH <= 33.0 else
            "Зубоальвеолярное укорочение моляров на верхней (U6-PP) и нижней (L6-MP) челюстях" if 25.0 <= UADH <= 31.0 and 38.0 <= UAHH <= 42.0 and UPDH < 21.0 and LPDH < 29.0 else
            "Зубоальвеолярная высота в норме на верхней челюсти в области резцов и на нижней челюсти в области резцов." if 25.0 <= UADH <= 31.0 and 38.0 <= UAHH <= 42.0 and UPDH < 21.0 and 29.0 <= LPDH <= 33.0 else
            "Зубоальвеолярная высота в норме на верхней челюсти в области резцов и на нижней челюсти в области моляров." if 25.0 <= UADH <= 31.0 and 38.0 <= UAHH <= 42.0 and 21.0 <= UPDH <= 24.0 and LPDH < 29.0 else
            "Зубоальвеолярное укорочение в области резцов на нижней (L1-MP) челюсти, зубоальвеолярное укорочение моляров на нижней (L6-MP) челюсти" if UADH > 31.0 and UAHH < 38.0 and 21.0 <= UPDH <= 24.0 and LPDH < 29.0 else
            "Зубоальвеолярное укорочение в области резцов на нижней (L1-MP) челюсти" if UADH > 31.0 and UAHH < 38.0 and 21.0 <= UPDH <= 24.0 and 29.0 <= LPDH <= 33.0 else
            "Зубоальвеолярное укорочение моляров на верхней (U6-PP) и нижней (L6-MP) челюстях" if UADH > 31.0 and 38.0 <= UAHH <= 42.0 and UPDH < 21.0 and LPDH < 29.0 else
            "Зубоальвеолярная высота в норме на верхней челюсти в области резцов и на нижней челюсти в области моляров." if UADH > 31.0 and 38.0 <= UAHH <= 42.0 and UPDH < 21.0 and 29.0 <= LPDH <= 33.0 else
            "Зубоальвеолярная высота в норме на верхней челюсти в области резцов и на нижней челюсти в области моляров." if UADH > 31.0 and 38.0 <= UAHH <= 42.0 and 21.0 <= UPDH <= 24.0 and LPDH < 29.0 else
            "Данные значения находятся за пределами ожидаемых диапазонов."  # Значение по умолчанию
    )
    else:
        uagh_uahh_updn_lpdn_result = "Отсутствуют необходимые данные для оценки."


    # Определение класса по значению G_SN_Po
    # G_SN_Po = float(values['G_SN_Po'][0]) if values['G_SN_Po'][0] else None
    # if G_SN_Po is not None:
    #     g_sn_po_class = (
    #         "в норме" if 10.0 <= G_SN_Po <= 14.0 else
    #         f"выпуклый на {G_SN_Po - 14.0:.1f}˚" if G_SN_Po > 14.0 else
    #         f"прямой на {10.0 - G_SN_Po:.1f}˚" if 2.0 <= G_SN_Po < 10.0 else
    #         f"вогнутый на {2.0 - G_SN_Po:.1f}˚" if G_SN_Po < 2.0 else
    #         ""  # Это условие не требуется, так как все случаи уже охвачены
    # )
    # else:
    #     g_sn_po_class = 'значение не найдено'

    # Преобразование значения G_SN_Po в float с проверкой на None
    G_SN_Po = float(values['G_SN_Po'][0]) if values.get('G_SN_Po') and values['G_SN_Po'][0] else None

    # Определение метки для градусов
    if G_SN_Po is not None:
        if 10.0 <= G_SN_Po <= 14.0:
            g_sn_po_class = "в норме"
            degree_info = ""
        elif G_SN_Po > 14.0:
            g_sn_po_class = "выпуклый"
            degree_info = f"на {G_SN_Po - 14.0:.1f}˚"
        elif 2.0 <= G_SN_Po < 10.0:
            g_sn_po_class = "прямой"
            degree_info = f"на {10.0 - G_SN_Po:.1f}˚"
        elif G_SN_Po < 2.0:
            g_sn_po_class = "вогнутый"
            degree_info = f"на {2.0 - G_SN_Po:.1f}˚"
        else:
            g_sn_po_class = ""
            degree_info = ""

    # Объединение основной метки и метки градусов
        # result = f"{g_sn_po_class} {degree_info}".strip()
    else:
        result = 'значение не найдено'



    # Преобразование значения Col_Sn_UL в float с проверкой на None
    Col_Sn_UL = float(values['Col_Sn_UL'][0]) if values.get('Col_Sn_UL') and values['Col_Sn_UL'][0] else None

    # Определение метки для градусов
    if Col_Sn_UL is not None:
        if 94.0 <= Col_Sn_UL <= 110.0:
            col_sn_ul_class = "в норме"
            degree_info = ""
        elif Col_Sn_UL > 110.0:
            col_sn_ul_class = "увеличению"
            degree_info = f"на {Col_Sn_UL - 110.0:.1f}˚"
        elif Col_Sn_UL < 94.0:
            col_sn_ul_class = "уменьшению"
            degree_info = f"на {94.0 - Col_Sn_UL:.1f}˚"
        else:
            col_sn_ul_class = ""
            degree_info = ""

    # Объединение основной метки и метки градусов
        # result = f"{col_sn_ul_class} {degree_info}".strip()
    else:
        result = 'значение не найдено'



    filled_text = template_text.format(
        ANB=values['ANB'][0] if values['ANB'][0] else 'значение не найдено',
        ANB_class=anb_class,
        Beta_Angle=values['Beta_Angle'][0] if values['Beta_Angle'][0] else 'значение не найдено',
        Beta_Angle_class=beta_angle_class,
        Wits_Appraisal=values['Wits_Appraisal'][0] if values['Wits_Appraisal'][0] else 'значение не найдено',
        Wits_Appraisal_class=wits_appraisal_class,
        Wits_appraisal_woman_class=wits_appraisal_woman_class,
        Wits_Appraisal_class_disproportion=wits_appraisal_class_disproportion,
        B_to_A_Point_Arc=values['B_to_A_Point_Arc'][0] if values['B_to_A_Point_Arc'][0] else 'значение не найдено',
        B_to_A_Point_Arc_class=b_to_a_point_arc_class,
        B_to_A_Point_Arc_direction=b_to_a_point_arc_direction,
        APDI=values['APDI'][0] if values['APDI'][0] else 'значение не найдено',
        Apdi_class=apdi_class,
        SN=values['SN'][0] if values['SN'][0] else 'значение не найдено',
        SN_result=sn_result,
        PNS_A=values['PNS_A'][0] if values['PNS_A'][0] else 'значение не найдено',
        Pns_sn_value_result=pns_sn_value_result,
        SNA=values['SNA'][0] if values['SNA'][0] else 'значение не найдено',
        SNA_class=sna_class,
        SN_Palatal_Plane=values['SN_Palatal_Plane'][0] if values['SN_Palatal_Plane'][0] else 'значение не найдено',
        SN_palatal_plane_class=sn_palatal_plane_class,
        Go_me_result=go_me_result,
        Go_Me=values['Go_Me'][0] if values['Go_Me'][0] else 'значение не найдено',
        Go_me_value_result=go_me_value_result,
        SNB=values['SNB'][0] if values['SNB'][0] else 'значение не найдено',
        SNB_class=snb_class,
        MP_SN=values['MP_SN'][0] if values['MP_SN'][0] else 'значение не найдено',
        MP_SN_class=mp_sn_class,
        Ar_Go_Me=values['Ar_Go_Me'][0] if values['Ar_Go_Me'][0] else 'значение не найдено',
        Ar_Go_Me_class=ar_go_me_class,
        Ar_Go_Na=values['Ar_Go_Na'][0] if values['Ar_Go_Na'][0] else 'значение не найдено',
        Ar_Go_Na_class=ar_go_na_class,
        Na_Go_Me=values['Na_Go_Me'][0] if values['Na_Go_Me'][0] else 'значение не найдено',
        Na_Go_Me_class=na_go_me_class,
        N_ANS=values['N_ANS'][0] if values['N_ANS'][0] else 'значение не найдено',
        ANS_Gn=values['ANS_Gn'][0] if values['ANS_Gn'][0] else 'значение не найдено',
        ANS_result=ans_result,
        ANS_result_class=ans_result_class,
        SGo_NGn=values['SGo_NGn'][0] if values['SGo_NGn'][0] else 'значение не найдено',
        Sgo_ngn_class=sgo_ngn_class,
        ANS_Xi_Pm=values['ANS_Xi_Pm'][0] if values['ANS_Xi_Pm'][0] else 'значение не найдено',
        FMA=values['FMA'][0] if values['FMA'][0] else 'значение не найдено',
        NaBa_PtGn=values['NaBa_PtGn'][0] if values['NaBa_PtGn'][0] else 'значение не найдено',
        I_VP_result=i_vp_result,
        I_VP_result_class=i_vp_result_class,
        ODI=values['ODI'][0] if values['ODI'][0] else 'значение не найдено',
        ODI_class=odi_class,
        U1_L1=values['U1_L1'][0] if values['U1_L1'][0] else 'значение не найдено',
        U1_L1_class=u1_l1_class,
        # Palatal_Plane=values['Palatal_Plane'][0] if values['Palatal_Plane'][0] else 'значение не найдено',
        Palatal_Plane_class=palatal_plane_class,
        IMPA_class=impa_class,
        Uagh_Uahh_Updn_Lpdn_result=uagh_uahh_updn_lpdn_result,
        # UADH=values['UADH'][0] if values['UADH'][0] else 'значение не найдено',
        # UAHH=values['UAHH'][0] if values['UAHH'][0] else 'значение не найдено',
        # UPDH=values['UPDH'][0] if values['UPDH'][0] else 'значение не найдено',
        # LPDH=values['LPDH'][0] if values['LPDH'][0] else 'значение не найдено',
        # G_SN_Po=values['G_SN_Po'][0] if values['G_SN_Po'][0] else 'значение не найдено',
        G_SN_PO_class=g_sn_po_class,
        Col_Sn_UL=values['Col_Sn_UL'][0] if values['Col_Sn_UL'][0] else 'значение не найдено',
        Сol_sn_ul_class=col_sn_ul_class,
        Upper_lip=values['Upper_lip'][0] if values['Upper_lip'][0] else 'значение не найдено',
        Lower_lip=values['Lower_lip'][0] if values['Lower_lip'][0] else 'значение не найдено',
    )

    doc.add_paragraph(filled_text)

    # doc.save(output_filename)
    # Создание документа
    doc = BytesIO()
    doc_obj = Document()
    doc_obj.add_paragraph(filled_text)
    doc_obj.save(doc)
    doc.seek(0)
    return doc

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        if 'file' not in request.files:
            return render_template('index.html', message="Файл не найден.")

        file = request.files['file']
        if file.filename == '':
            return render_template('index.html', message="Файл не выбран.")

        gender = request.form.get('gender', 'мужчина').strip().lower()
        if not gender in ['мужчина', 'женщина']:
            return render_template('index.html', message="Неверно указан пол.")

        file_content = file.read()
        patterns = {
        'ANB': r'ANB\s*\(.*?\)\s*([+-]?\d*\.?\d+)',
        'Beta_Angle': r'Beta Angle\s*\(.*?\)\s*([+-]?\d*\.?\d+)',
        'Wits_Appraisal': r'Wits Appraisal\s*\(.*?\)\s*([+-]?\d*\.?\d+)',
        'B_to_A_Point_Arc': r'B\s*to\s*A\s*Point\s*Arc\s+([+-]?\d*\.?\d+)',
        'APDI': r'Anteroposterior Dysplasia \(APDI\)\s+([+-]?\d*\.?\d+)',
        'SN': r'Anterior Cranial Base \(SN\) \(mm\)\s*([+-]?\d*\.?\d+)',
        'PNS_A': r'PNS-A\s*\(mm\)\s*([+-]?\d+\.?\d*)',
        'SNA': r'SNA\s*\(є\)\s*(-?\d+\.\d+|-?\d+)',
        'SN_Palatal_Plane': r"Cranio-Mx Base/SN-Palatal Plane\s*\(є\)\s*([-+]?\d*\.\d+|\d+)",
        'Go_Me': r"Mandibular Body Length \(Go-Me\) \(mm\)\s*(-?\d+\.\d+)",
        'SNB': r'SNB\s+\(є\)\s+([+-]?\d*\.\d+|\d+)',
        'MP_SN': r'MP\s*-\s*SN\s*\(є\)\s*(-?\d+\.\d+)',
        'Ar_Go_Me': r'Gonial/Jaw Angle \(Ar-Go-Me\) \(є\)\s*(-?\d+(\.\d+)?)',
        'Ar_Go_Na': r'Upper Gonial Angle \(Ar-Go-Na\) \(є\)\s*(-?\d+(\.\d+)?)',
        'Na_Go_Me': r'Lower Gonial Angle \(Na-Go-Me\) \(є\)\s*(-?\d+(\.\d+)?)',
        'N_ANS': r'Upper Face Height \(N-ANS\) \(mm\)\s*(-?\d+\.\d+|\d+)',
        'ANS_Gn': r'Lower Face Height \(ANS-Gn\) \(mm\)\s*(-?\d+\.\d+|\d+)',
        'SGo_NGn': r'SGo/NGn \(%\)\s+(-?\d+(\.\d+)?)',
        'ANS_Xi_Pm': r'Lower Face Height \(ANS\-Xi\-Pm\)\(є\)\s+(-?\d+(\.\d+)?)',
        'FMA': r'FMA \(MP\-FH\) \(є\)\s+(-?\d+(\.\d+)?)',
        'NaBa_PtGn': r'Facial Axis\-Ricketts \(NaBa\-PtGn\)\(є\)\s+(-?\d+(\.\d+)?)',
        'ODI': r'Overbite Depth Indicator \(ODI\)\s+([+-]?\d*\.\d+|\d+)',
        'U1_L1': r'Interincisal Angle \(U1-L1\) \(є\)\s+([+-]?\d*\.\d+|\d+)',
        'Palatal_Plane': r'U1\s*-\s*Palatal\s*Plane\s*\(є\)\s+([+-]?\d*\.\d+|\d+)',
        'IMPA': r'IMPA \(L1-MP\)\s+\(є\)\s+([+-]?\d*\.\d+|\d+)',
        'UADH': r'U1\s*-\s*PP\s*\(UADH\)\s*\(mm\)\s+([+-]?\d*\.\d+|\d+)',
        'UAHH': r'L1\s*-\s*MP\s*\(LADH\)\s*\(mm\)\s+([+-]?\d*\.\d+|\d+)',
        'UPDH': r'U6\s*-\s*PP\s*\(UPDH\)\s*\(mm\)\s+([+-]?\d*\.\d+|\d+)',
        'LPDH': r'L6\s*-\s*MP\s*\(LPDH\)\s*\(mm\)\s+([+-]?\d*\.\d+|\d+)',
        'G_SN_Po': r'Facial\s*Convexity\s*\(G\'-Sn-Po\'\)\s*\(є\)\s+([+-]?\d*\.\d+|\d+)',
        'Col_Sn_UL': r'Nasolabial\s*Angle\s*\(Col-Sn-UL\)\s*\(є\)\s+([+-]?\d*\.\d+|\d+)',
        'Upper_lip': r'Upper\s*Lip\s*to\s*E-Plane\s*\(mm\)\s+([+-]?\d*\.\d+|\d+)',
        'Lower_lip': r'Lower\s*Lip\s*to\s*E-Plane\s*\(mm\)\s+([+-]?\d*\.\d+|\d+)',
        }
        values = extract_values_from_txt(file_content, patterns)
        doc = create_document_with_values(values, gender)
        return send_file(doc, as_attachment=True, download_name="report.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    return render_template('index.html')

if __name__ == "__main__":
    app.run(debug=True)